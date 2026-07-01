#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génération du rapport Planning de Réalisation en format Word (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import shutil
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGE_SRC = r"C:\Users\THINKPAD\.gemini\antigravity\brain\b57a3f08-c7bc-42f3-8b87-54b633bb756a\media__1782631141613.png"
IMAGE_LOCAL = os.path.join(SCRIPT_DIR, "gantt_capture.png")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "Rapport_Planning_Realisation.docx")

# Copier l'image
if not os.path.exists(IMAGE_LOCAL):
    shutil.copy2(IMAGE_SRC, IMAGE_LOCAL)


def set_cell_shading(cell, color):
    """Applique une couleur de fond à une cellule."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC"):
    """Ajoute des bordures fines à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_paragraph(doc, text, size=11, bold=False, italic=False, color=None, alignment=None, space_after=6):
    """Ajoute un paragraphe formaté."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p, run


def create_cover_page(doc):
    """Crée la page de garde."""

    # Espacement haut
    for _ in range(3):
        add_paragraph(doc, '', size=12, space_after=12)

    # Icône
    add_paragraph(doc, '🏗️', size=36, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Titre principal
    add_paragraph(doc, 'PLANNING DE RÉALISATION', size=28, bold=True,
                  color=(15, 23, 42), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # Ligne décorative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(59, 130, 246)
    p.paragraph_format.space_after = Pt(4)

    # Sous-titre
    add_paragraph(doc, "Réalisation des travaux d'un pôle médical", size=14,
                  color=(100, 116, 139), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # Tableau d'informations du projet
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    infos = [
        ('Objet', 'Planning de réalisation — Pôle médical'),
        ('Délai d\'exécution', 'À partir de l\'ODS notifié au service co-contractant par le service contractant'),
        ('Version 1', 'Délai de 24 mois'),
        ('Version 2', 'Délai optimisé de 16 mois (facteur ×0.667)'),
        ('Date du document', datetime.now().strftime('%d/%m/%Y')),
        ('Statut', 'Document confidentiel — Usage interne'),
    ]

    for i, (label, value) in enumerate(infos):
        # Colonne label
        cell_l = table.cell(i, 0)
        cell_l.width = Cm(5)
        p = cell_l.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(59, 130, 246)
        set_cell_shading(cell_l, "F0F5FF")
        set_cell_borders(cell_l, "D0DCEE")

        # Colonne valeur
        cell_v = table.cell(i, 1)
        cell_v.width = Cm(10)
        p = cell_v.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(51, 65, 85)
        set_cell_borders(cell_v, "D0DCEE")

    # Espace
    add_paragraph(doc, '', size=12, space_after=40)

    # Badge confidentiel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('📋  DOCUMENT CONFIDENTIEL — USAGE INTERNE')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(59, 130, 246)
    p.paragraph_format.space_after = Pt(30)

    # Date de génération
    add_paragraph(doc, f'Généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}',
                  size=9, italic=True, color=(150, 150, 150),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    # Saut de page
    doc.add_page_break()


def create_gantt_page(doc):
    """Page avec le diagramme de Gantt."""

    # Titre
    add_paragraph(doc, 'Diagramme de Gantt — Version 16 Mois', size=18, bold=True,
                  color=(15, 23, 42), space_after=8)

    # Description
    add_paragraph(doc,
        'Le diagramme ci-dessous présente le planning de réalisation recalculé sur 16 mois '
        '(facteur de compression ×0.667 appliqué à toutes les durées et dates de début '
        'par rapport à la version originale de 24 mois).',
        size=10, color=(100, 116, 139), space_after=10)

    # Encadré délai
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F0F8FF")
    set_cell_borders(cell, "3B82F6")
    p = cell.paragraphs[0]
    run = p.add_run("Délai d'exécution : ")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(59, 130, 246)
    run = p.add_run("À partir de l'ODS notifié au service co-contractant par le service contractant.")
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(51, 65, 85)

    add_paragraph(doc, '', size=6, space_after=8)

    # Image du Gantt
    doc.add_picture(IMAGE_LOCAL, width=Inches(6.3))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Légende
    add_paragraph(doc, 'Figure 1 — Diagramme de Gantt interactif, version 16 mois (capture écran)',
                  size=8, italic=True, color=(120, 120, 120),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    doc.add_page_break()


def create_recap_page(doc):
    """Page récapitulatif avec tableau comparatif."""

    # Titre
    add_paragraph(doc, 'Récapitulatif des tâches', size=18, bold=True,
                  color=(15, 23, 42), space_after=6)

    add_paragraph(doc, 'Comparaison des durées entre la version 24 mois et la version optimisée 16 mois.',
                  size=10, color=(100, 116, 139), space_after=12)

    # Tableau comparatif
    headers = ['#', 'Tâche', 'V1 (24 mois)', '→', 'V2 (16 mois)', 'Début V2']
    tasks = [
        ('1', 'Installation de chantier', '15 jours', '→', '10 jours', 'M0'),
        ('2', 'Terrassements / Excavations', '3 mois', '→', '2 mois', '10J'),
        ('3', 'Infrastructure', '7 mois', '→', '5 mois', 'M2'),
        ('4', 'Superstructure', '12 mois', '→', '8 mois', 'M5'),
        ('5', 'Maçonnerie et enduits', '12 mois', '→', '8 mois', 'M5'),
        ('6', 'VRD', '17 mois', '→', '11 mois', 'M2½'),
        ('7', 'Électricité / Plomberie', '17 mois', '→', '11 mois', 'M3'),
        ('8', 'Clim / Chauffage central', '16 mois', '→', '11 mois', 'M3'),
        ('9', 'Finition sol / Peinture', '13 mois', '→', '9 mois', 'M6'),
        ('10', 'Nettoyage', '2 mois', '→', '1 mois ½', 'M14'),
        ('11', 'Essais + Formation', '2 mois ½', '→', '1 mois ½', 'M14'),
        ('12', 'Réception provisoire', 'Mois 24', '→', 'Mois 16', 'M16'),
    ]

    table = doc.add_table(rows=1 + len(tasks), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    header_colors = {
        0: "0F172A", 1: "0F172A", 2: "1E3A5F", 3: "0F172A", 4: "7C2D12", 5: "7C2D12"
    }
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_colors.get(i, "0F172A"))
        set_cell_borders(cell, "2A3550")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Données
    for row_idx, row_data in enumerate(tasks):
        bg = "F5F7FA" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            set_cell_shading(cell, bg)
            set_cell_borders(cell, "DEE2E6")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

            # Couleurs par colonne
            if col_idx == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42)
            elif col_idx == 2:
                run.font.color.rgb = RGBColor(59, 130, 246)
            elif col_idx == 3:
                run.font.color.rgb = RGBColor(245, 158, 11)
                run.font.size = Pt(12)
            elif col_idx == 4:
                run.font.bold = True
                run.font.color.rgb = RGBColor(249, 115, 22)
            elif col_idx == 5:
                run.font.color.rgb = RGBColor(100, 116, 139)
            else:
                run.font.color.rgb = RGBColor(51, 65, 85)

    add_paragraph(doc, '', size=6, space_after=14)

    # Note
    note_table = doc.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = note_table.cell(0, 0)
    set_cell_shading(cell, "FFFBEB")
    set_cell_borders(cell, "F59E0B")
    p = cell.paragraphs[0]
    run = p.add_run("Note : ")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(180, 83, 9)
    run = p.add_run("Le facteur de compression ×0.667 (16/24) a été appliqué uniformément à toutes les durées et dates de début.")
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(120, 80, 20)


def main():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    print("[...] Page de garde...")
    create_cover_page(doc)

    print("[...] Page Gantt...")
    create_gantt_page(doc)

    print("[...] Page récapitulatif...")
    create_recap_page(doc)

    doc.save(OUTPUT_DOCX)
    print(f"\n[OK] Rapport Word généré : {OUTPUT_DOCX}")


if __name__ == '__main__':
    main()
